# coding=utf-8
import requests
from bs4 import BeautifulSoup
import docx
import openpyxl


# crawling
print "Getting list of articles..."
response = requests.get("http://news.naver.com/main/list.nhn?sid1=105&mid=sec&mode=LSD&date=20160127&page=1")
response.raise_for_status()

# a tag parse
soup = BeautifulSoup(response.text, "html.parser")
art_list = soup.select('.type06_headline > li > dl > dt > a')
art_list.extend(soup.select('.type06 > li > dl > dt > a'))

# sinmunsa parse
sinmunsa_dict = {}
sinmunsa_list = soup.select('.writing')
for sinmunsa in sinmunsa_list:
    sinmunsa_dict.setdefault(sinmunsa.get_text(), 0)
    sinmunsa_dict[sinmunsa.get_text()] += 1


# excel
print "Writing sinmunsa to excel..."
wb = openpyxl.Workbook()
sheet = wb.get_sheet_by_name('Sheet')
sheet['A1'] = '신문사 이름'
sheet['B1'] = '작성된 기사의 수'

cell_idx = 2
for k, v in sinmunsa_dict.items():
    sheet.cell(row=cell_idx, column=1).value = k
    sheet.cell(row=cell_idx, column=2).value = v
    cell_idx += 1

wb.save('sinmunsa.xlsx')


# href url into list
real_list = []
for art in art_list:
    real_list.append(art['href'])


# doc writing
doc = docx.Document()
print "Writing articles to docs..."
for i in set(real_list):
    response = requests.get(i)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")

    title = soup.select('#articleTitle')[0].get_text()
    content = soup.select('#articleBodyContents')[0].get_text()

    doc.add_heading(title, 0)
    doc.add_paragraph(content)

    doc.add_page_break()

print "Now saving..."
doc.save('articles.docx')
print "Done"