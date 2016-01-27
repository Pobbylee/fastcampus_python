# coding=utf-8
import requests
from bs4 import BeautifulSoup
import docx

# crawling
response = requests.get("http://news.naver.com/main/read.nhn?mode=LSD&mid=sec&sid1=105&oid=029&aid=0002327059")
response.raise_for_status()

soup = BeautifulSoup(response.text, "html.parser")

title = soup.select('#articleTitle')[0].get_text()
content = soup.select('#articleBodyContents')[0].get_text()

# doc writing
doc = docx.Document()

doc.add_heading(title, 0)
doc.add_paragraph(content)

doc.add_page_break()

doc.add_heading(title, 0)
doc.add_paragraph(content)

doc.add_page_break()

doc.save('article.docx')