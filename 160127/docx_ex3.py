import docx
doc = docx.Document()

doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.add_heading('Header 5', 5)
doc.add_heading('Header 6', 6)
doc.add_heading('Header 7', 7)
doc.add_heading('Header 8', 8)

# for i in range(0, 8):
#     doc.add_heading('Header' + str(i), i)

doc.save('headings.docx')