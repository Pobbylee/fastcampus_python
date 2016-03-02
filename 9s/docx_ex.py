import docx
doc = docx.Document('demo.docx')
# print doc.paragraphs
# print len(doc.paragraphs)

# print doc.paragraphs[0].text
#
# print doc.paragraphs[1].text

# for i in doc.paragraphs:
#     print i.text
#

# print doc.paragraphs[1].runs
# print len(doc.paragraphs[1].runs)

print doc.paragraphs[1].runs[0].text

print doc.paragraphs[1].runs[1].text

print doc.paragraphs[1].runs[2].text

print doc.paragraphs[1].runs[3].text